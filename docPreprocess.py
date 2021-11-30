import os.path
from tqdm import tqdm
import docx
import pandas as pd
import json
import re
from elementExtract.ruleExt1 import RuleExtract1
from copy import deepcopy
import argparse


def get_json_from_doc(filename, output="data.json"):
    """
    提取doc文件中的表格，每行转为一个item，包含多个属性，不分层级
    :param filename: 待处理的doc文件路径
    :param output: 存储的json文件路径
    :return:
    """
    if not os.path.exists(filename):
        print(f"文件{filename}不存在，请先检查！")
        return
    if os.path.exists(output):
        print(f"文件{output}已存在！")
        return
    doc = docx.Document(filename)
    tables = doc.tables
    assert len(tables) != 0
    table = tables[0]
    rows = table.rows
    assert len(rows)
    names = rows[0].cells
    out = open(output, 'w')
    results = []
    for i, row in enumerate(tqdm(rows)):
        if i == 0:
            continue
        item = {}
        for j, cell in enumerate(row.cells):
            item[names[j].text] = cell.text
        if item.get("序号", None):
            item["序号"] = i
        results.append(item)
    json.dump(results, out, ensure_ascii=True, indent=4)


def table2json(input_path):
    """
    提取文件中的表格，每行转为一个item，包含多个属性，不分层级
    :param input_path: 文件路径
    :return: json文件
    """
    dicts = []
    if input_path[-5:] == '.docx':
        f = docx.Document(input_path)
        for table in f.tables:
            attributes = [x.text.replace('\n', '').replace(' ', '') for x in table.rows[0].cells]
            add_attr = []
            for row in tqdm(table.rows[1:]):
                text_set = list(set([x.text for x in row.cells]))
                if len(text_set) == 2:
                    tmp = row.cells[0].text.replace('\n', '').replace(' ', '')
                    if ('（' in tmp or '(' in tmp) and len(add_attr) == 2:
                        add_attr.pop()
                    if tmp in ['一', '二', '三', '四', '五', '六', '七', '八', '九', '零', '〇']:
                        add_attr = []
                    add_attr.append(row.cells[1].text.replace(' ', '').replace('\n', ''))
                else:
                    item = {'类型': add_attr.copy()}
                    for idx, att in enumerate(attributes):
                        item[str(att)] = row.cells[idx].text
                        # 日期统一用 Y-m-d形式记录
                        if '日期' in att:
                            item[str(att)] = item[str(att)].replace('/', '-')
                    dicts.append(item)
    elif input_path[-5:] == '.xlsx':
        df = pd.read_excel(input_path, header=1)
        attributes = [x.replace('\n', '').replace(' ', '') for x in df.columns]
        add_attr = []
        for index, row in df.iterrows():
            if pd.isna(row[4]):  # 标准编号不能为空
                tmp = row[0].replace('\n', '').replace(' ', '')
                if ('（' in tmp or '(' in tmp) and len(add_attr) == 2:
                    add_attr.pop()
                if tmp in ['一', '二', '三', '四', '五', '六', '七', '八', '九', '零', '〇']:
                    add_attr = []
                add_attr.append(row[1].replace(' ', '').replace('\n', ''))
            else:
                item = {'类型': add_attr.copy()}
                for attr in attributes:
                    item[str(attr)] = str(row[str(attr)])
                    if '日期' in str(attr):
                        item[str(attr)] = item[str(attr)].split()[0]
                dicts.append(item)

    out_path = os.path.splitext(input_path)[0] + '.json'
    with open(out_path, 'w', encoding='utf-8') as f:
        json.dump(dicts, f, ensure_ascii=False, indent=2)


def clear_json(input_path):
    """
    对json文件数据清洗，去除多余换行符，序号补全，统一标准编号，去除编号空格
    :param input_path:
    :return:
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    for i, item in enumerate(data):
        keys = list(item.keys())
        for key in keys:
            if key == '序号':
                item['序号'] = str(i + 1)
            elif key == '相关条款说明' or key == '最新实施/新修法律法规解读' or key == '检查内容':
                reg = r'\n'
                item[key] = re.sub(reg, '', item[key])
            elif key == '标准编号' or key == '文件编号':
                item[key] = item[key].replace('\n', '').replace(' ', '')
                item['标准编号'] = item[key]  # 统一添加标准编号
            elif key != '类型':
                item[key] = item[key].replace('\n', '')
        if item.get('文件编号', True) is not True:
            del item['文件编号']
        data[i] = item
    with open(input_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def form_data(input_path):
    """
    data.json文件提出“标准编号”属性（去除空格），可能有多个编号，故按list存储
    :param input_path:
    :return:
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    new_data = []
    for i, item in enumerate(data):
        item['检查依据'] = item['检查依据'].replace(' ', '')
        reg = '(《.*?》)([^条]*)第(.*?)条'
        finds = re.findall(reg, item['检查依据'])
        # for test
        # if i in [44, 473, 692, 1051]:
        #     print(">>>>>>>>")
        for find in finds:
            new_item = deepcopy(item)
            new_item['检查依据'] = find[0]
            new_item['标准编号'] = find[1]
            new_item['条目'] = find[2]
            new_data.append(new_item)
    with open(input_path, 'w', encoding='utf-8') as f:
        json.dump(new_data, f, ensure_ascii=False, indent=4)


def extract_rule_data(input_path, output_path):
    """
    对data.json文件中的“检查内容”进行解析，得到展示形式和应用形式的规则
    :param input_path:
    :param output_path:
    :return:
    """
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    ruleExtr1 = RuleExtract1()
    for item in tqdm(data):
        ruleExtr1.setSent(item['检查内容'])
        ruleExtr1.parser()
        item['viewForm'] = ruleExtr1.genViewVer()
        item['appForm'] = ruleExtr1.genAppVer()
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# 解析rule.json中viewForm和appForm属性
def analysis_rule(input_path):
    with open(input_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    for item in data:
        viewForm = item['viewForm'][0]
        for key in viewForm.keys():
            print(key)
            _key = viewForm[key]


# TODO: 将2020和2021文件合并
def json_merge(out_path, **input_path):
    pass


def arg_parser():
    parser = argparse.ArgumentParser(description="doc preprocess")
    parser.add_argument("--input", default="data/隐患排查大全附标准依据.docx", type=str, help="input file path")
    parser.add_argument("--doc2json", default=False, action="store_true", help="from doc to json")
    parser.add_argument("--clear", default=False, action="store_true", help="clear json and form data")
    parser.add_argument("--extract", default=False, action="store_true", help="extract rule data")

    args = parser.parse_args()
    return args


def main(args):
    output_json = "data/data.json"
    if not os.path.exists("data"):
        os.makedirs("data")
    # 第一步： 从doc到json文件——doc处理速度很慢，json读取较快
    if args.doc2json:
        print("start get_json_from_doc ...")
        get_json_from_doc(args.input, output_json)

    if args.clear:
        print("start clear ...")
        # 第二步：对提取出的json文件做清洗，删除多余空格等
        clear_json(output_json)
        # 第三步：处理检查依据，抽取标准编号作为唯一标志符
        form_data(output_json)

    # 第四步；从content中抽取出对应内容
    if args.extract:
        print("start extract ...")
        extract_rule_data('data/data.json', 'data/rule.json')


if __name__ == '__main__':
    # main(arg_parser())
    # pass
    # table2json('data/law/安全生产法律法规列表.xlsx')
    # table2json('data/law/2020最新安全生产法律法规清单.docx')
    # table2json('data/law/安全法规列表-2.docx')
    clear_json('data/law/安全生产法律法规列表.json')
    clear_json('data/law/安全法规列表-2.json')
    clear_json('data/law/2020最新安全生产法律法规清单.json')
    # clearJson('data/data.json')
